from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_paragraph(doc, text=''):
    return doc.add_paragraph(text)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    return table

# Title
title = doc.add_heading('Guías DAS – Business Rules (MLB)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 1. Overview
add_heading(doc, '1. Overview and Scope')
add_paragraph(doc, 'Guías DAS is a feature within Mercado Pago that allows MEI merchants (Microempreendedor Individual) registered in Brazil to emit and pay their monthly DAS (Documento de Arrecadação do Simples Nacional) directly from the platform.')
add_paragraph(doc)
add_paragraph(doc, 'This document explains:')
for item in [
    'Who is eligible to use Guías DAS',
    'What the user can do within the flow',
    'How guide states work and how navigation between periods operates',
    'What users should do when they need to pay a specific month',
    'Common issues and how to handle them',
]:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# 2. Value Proposition
add_heading(doc, '2. Value Proposition')
add_paragraph(doc, 'Guías DAS helps MEI merchants:')
for item in [
    'Emit their monthly DAS guide for any desired fiscal period',
    'Pay the emitted guide directly from Mercado Pago, without leaving the platform',
    'View the payment status of each period in a centralized way',
    'Avoid using external government portals for DAS emission and payment',
]:
    doc.add_paragraph(item, style='List Bullet')
add_paragraph(doc)
add_paragraph(doc, 'The main goal is to simplify MEI tax compliance by integrating the DAS emission and payment flow into a single experience within Mercado Pago.')

doc.add_paragraph()

# 2.5 User Journey
add_heading(doc, '3. How It Works — Step-by-Step User Journey')
add_paragraph(doc, 'The complete flow for a MEI user to emit and pay their DAS through Mercado Pago is:')
doc.add_paragraph()

steps = [
    ('Step 1 — Access Central de Pagos',
     'The user logs into Mercado Pago and navigates to Central de Pagos (Payment Center). This is the hub for all tax and payment obligations.'),
    ('Step 2 — Enter Guías DAS',
     'Within Central de Pagos, the user selects the Guías DAS section. This section is only visible and accessible to users with a MEI-registered account.'),
    ('Step 3 — Select the fiscal period',
     'The default view shows the current month\'s DAS. If the user needs to pay a different month, they navigate to:\n  • Impuestos vencidos — for overdue/past months\n  • Impuestos futuros — for upcoming months'),
    ('Step 4 — Emit the boleto',
     'The user clicks the action to emit the DAS guide for the selected period. The system generates the boleto (DAS document) with the correct amount and due date for that fiscal period.'),
    ('Step 5 — Pay directly from Mercado Pago',
     'The user pays the emitted boleto directly within the platform, without needing to leave Mercado Pago or access any government portal. Payment is processed immediately.'),
    ('Step 6 — Confirmation',
     'The guide status updates to Paid. The government portal (PGMEI / Portal do Simples Nacional) may take up to 48 hours to reflect the payment — this is expected behavior.'),
]

for title_step, desc in steps:
    p = doc.add_paragraph()
    r = p.add_run(title_step)
    r.bold = True
    p.add_run('\n' + desc)

doc.add_paragraph()
p = add_paragraph(doc, 'Key rule: ')
p.runs[0].bold = True
p.add_run('The user does NOT need to use any external government site to emit or pay their DAS. The entire flow — from emission to payment — happens inside Mercado Pago, starting from Central de Pagos.')

doc.add_paragraph()

# 3. Eligibility
add_heading(doc, '4. Eligibility Requirements')
add_paragraph(doc, 'To access the Guías DAS flow, the user must meet all of the following conditions:')
doc.add_paragraph()
add_table(doc,
    ['Requirement', 'Detail'],
    [
        ['Market', 'MLB (Brazil) only'],
        ['Account type', 'Must be a MEI (Microempreendedor Individual) registered user'],
        ['Mercado Pago account', 'No specific account type required beyond MEI registration'],
    ]
)
doc.add_paragraph()
p = add_paragraph(doc, 'Important: ')
p.runs[0].bold = True
p.add_run('If the user is not a MEI (e.g., ME, EPP, or other company types), they cannot access the Guías DAS flow. In those cases, they must be redirected to the ')
r = p.add_run('boleto de pago')
r.bold = True
p.add_run(' flow, where they can enter a barcode manually to pay their tax obligations.')

doc.add_paragraph()

# 4. Functionalities
add_heading(doc, '5. Included Functionalities')
add_table(doc,
    ['Functionality', 'Description'],
    [
        ['DAS Emission', 'Emit a DAS guide for the desired fiscal period'],
        ['In-platform Payment', 'Pay the emitted guide directly from Mercado Pago'],
        ['Payment Status Visualization', 'See whether a given month has been paid or not'],
        ['Multi-period Navigation', 'Access overdue taxes and future taxes from dedicated sections'],
    ]
)

doc.add_paragraph()

# 5. Guide States
add_heading(doc, '6. Guide States')
add_paragraph(doc, 'Each monthly DAS guide can be in one of two states:')
doc.add_paragraph()
add_table(doc,
    ['State', 'Description', 'Action Available'],
    [
        ['Not Paid', 'The DAS for that month has not been paid through any channel', 'Enables the user to emit and pay through Mercado Pago'],
        ['Paid', 'The DAS for that month was already paid (can be from an external site or Mercado Pago)', 'No payment action available; displayed as paid'],
    ]
)
doc.add_paragraph()
p = add_paragraph(doc, 'Important: ')
p.runs[0].bold = True
p.add_run('If a user paid their DAS through an external platform (e.g., gov.br or another bank), the guide will still show as paid. The system reflects the real payment status regardless of where the payment was made.')

doc.add_paragraph()

# 6. Navigation
add_heading(doc, '7. Navigating Between Fiscal Periods')
add_paragraph(doc, 'The Guías DAS flow displays one guide per month. To pay or emit a guide for a different period, the user must navigate accordingly:')
doc.add_paragraph()
add_table(doc,
    ['Scenario', 'Where to go'],
    [
        ["Current month's DAS", 'Default view in Guías DAS'],
        ['Previous unpaid months', '"Impuestos vencidos" (Overdue taxes) section'],
        ['Future months', '"Impuestos futuros" (Future taxes) section'],
    ]
)
doc.add_paragraph()
p = add_paragraph(doc, 'Golden Rule for Support (Bot and Agents): ')
p.runs[0].bold = True
p.add_run('If a user says they cannot find the guide for a specific month or that they want to pay a past/future month, they must be directed to the ')
r = p.add_run('Impuestos vencidos')
r.italic = True
p.add_run(' or ')
r2 = p.add_run('Impuestos futuros')
r2.italic = True
p.add_run(' sections, not to the default Guías DAS view.')

doc.add_paragraph()

# 7. Common Issues
add_heading(doc, '8. Common Issues and Resolution')
add_heading(doc, '8.1 Payment not reflected on the government portal', level=2)
add_paragraph(doc, 'Symptom: The user paid their DAS through Mercado Pago but the payment does not appear on the government site (e.g., Portal do Simples Nacional / PGMEI).')
doc.add_paragraph()
add_paragraph(doc, 'Cause: Government systems have a processing delay.')
doc.add_paragraph()
p = add_paragraph(doc, 'Resolution: ')
p.runs[0].bold = True
p.add_run('Inform the user that the payment was processed successfully on Mercado Pago\'s side. Government portals may take up to 48 hours to reflect the payment. No action is required — they should check again after that period.')
doc.add_paragraph()
p = add_paragraph(doc, 'Note: ')
p.runs[0].bold = True
p.add_run('This is the most common support inquiry for this product. Agents and bots must communicate this proactively before escalating.')

doc.add_paragraph()

# 8. Non-MEI
add_heading(doc, '9. Non-MEI Users — Redirection Rule')
add_paragraph(doc, 'If a user who is not a MEI attempts to access the Guías DAS flow or asks about paying DAS:')
for item in [
    'They cannot use this flow.',
    'They must be redirected to the boleto de pago feature, where they can enter a barcode to pay any tax obligation manually.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# 9. Links
add_heading(doc, '10. Useful Links (MLB – Brazil)')
add_paragraph(doc, 'Note: Links for desktop and mobile versions are being confirmed. This section will be updated once available.')
doc.add_paragraph()
add_table(doc,
    ['Destination', 'Desktop', 'Mobile'],
    [
        ['Guías DAS main flow', '(pending)', '(pending)'],
        ['Impuestos vencidos', '(pending)', '(pending)'],
        ['Impuestos futuros', '(pending)', '(pending)'],
        ['Boleto de pago (non-MEI)', '(pending)', '(pending)'],
    ]
)

doc.save('/Users/woviedoalist/Desktop/Guias_DAS_Business_Rules.docx')
print("Documento creado: /Users/woviedoalist/Desktop/Guias_DAS_Business_Rules.docx")
